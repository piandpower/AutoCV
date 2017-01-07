import docx
doc = docx.Document('demo.docx')
print(len(doc.paragraphs))
print(doc.paragraphs[0].text)
print(doc.paragraphs[1].text)
print(len(doc.paragraphs[1].runs)) #runs = plain text + bold/italic/styled text + more plain text = 3 in this case
allruns = doc.paragraphs[1].runs[0].text + doc.paragraphs[1].runs[1].text + doc.paragraphs[1].runs[2].text + doc.paragraphs[1].runs[3].text
print(allruns) #allruns is the same as doc.paragraphs[1].runs
